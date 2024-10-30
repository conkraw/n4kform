import io
import os
import base64
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from google.cloud import storage
from docx import Document

# Function to upload document to Firebase Storage
def upload_to_storage(bucket_name, file_name, file_data):
    storage_client = storage.Client()
    bucket = storage_client.bucket(bucket_name)
    
    blob = bucket.blob(file_name)
    blob.upload_from_file(file_data, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    return blob.public_url  # Return the public URL

# Function to send email with attachment
def send_email_with_attachment(to_email, subject, text, html, file_path):
    from_email = 'your_email@example.com'  # Replace with your email
    password = 'your_email_password'         # Replace with your password

    # Create a multipart email
    msg = MIMEMultipart()
    msg['From'] = from_email
    msg['To'] = to_email
    msg['Subject'] = subject

    # Attach the email body
    msg.attach(MIMEText(text, 'plain'))
    msg.attach(MIMEText(html, 'html'))

    # Attach the Word document
    with open(file_path, 'rb') as attachment:
        part = MIMEBase('application', 'octet-stream')
        part.set_payload(attachment.read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', f'attachment; filename={file_path}')
        msg.attach(part)

    # Send the email
    try:
        with smtplib.SMTP('smtp.example.com', 587) as server:  # Replace with your SMTP server
            server.starttls()
            server.login(from_email, password)
            server.send_message(msg)
            print("Email sent successfully!")
    except Exception as e:
        print(f"Error sending email: {e}")

# Main script
def main(text_input):
    # Create a Word document
    doc = Document()
    doc.add_heading('Text Input Submission', level=1)
    doc.add_paragraph(text_input)

    # Save the document to a local file
    file_name = 'text_input_submission.docx'
    doc.save(file_name)

    # Upload document to Firebase Storage (optional)
    bucket_name = 'your_bucket_name'  # Replace with your bucket name
    file_url = upload_to_storage(bucket_name, file_name, open(file_name, 'rb'))

    # Send email with the attachment
    send_email_with_attachment(
        to_email='ckrawiec@pennstatehealth.psu.edu',
        subject='Hello from Firebase!',
        text='This is the plaintext section of the email body.',
        html='This is the <code>HTML</code> section of the email body.',
        file_path=file_name  # Path to the local file
    )

    # Clean up the local file
    os.remove(file_name)

# Example usage
if __name__ == '__main__':
    text_input = "Your input text here."  # Replace with actual input
    main(text_input)

